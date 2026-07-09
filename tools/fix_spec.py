"""Final fix: use unicode normalization and exact heading indexing."""
import json, unicodedata
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"D:\课设\功能规格说明模板.docx"
OUT = r"D:\project-89\功能规格说明_冷戰兵棋推演系统.docx"
BASE = r"D:\project-89\tools"

doc = Document(SRC)

# Load data with normalized keys
data = {}
for fn in ["spec_content.json","spec_content_2.json","spec_content_3.json","spec_content_4.json"]:
    with open(f"{BASE}\\{fn}", "r", encoding="utf-8") as f:
        d = json.load(f)
        for k, v in d.items():
            data[unicodedata.normalize("NFC", k)] = v

# Find heading style IDs
h2_id = h3_id = None
for s in doc.styles:
    if s.name == "Heading 2": h2_id = s.style_id
    elif s.name == "Heading 3": h3_id = s.style_id

def mke(text, sid=None):
    p = OxmlElement("w:p")
    if sid:
        pr = OxmlElement("w:pPr"); ps = OxmlElement("w:pStyle")
        ps.set(qn("w:val"), sid); pr.append(ps); p.append(pr)
    r = OxmlElement("w:r"); t = OxmlElement("w:t")
    t.text = text; t.set(qn("xml:space"), "preserve"); r.append(t); p.append(r)
    return p

# First pass: normalize heading text and rename headings
# Build map of paragraph index to replacement content
replacements = {}
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    st = p.style.name
    nt = unicodedata.normalize("NFC", t)

    if st.startswith("Heading") and "xxx" in nt:
        if "3.2" in nt:
            p.clear(); p.add_run("3.2 战场网格与机动模块 (Battlefield Grid & Movement)")
        elif "3.3" in nt:
            p.clear(); p.add_run("3.3 战斗计算模块 (Combat Resolution)")

    # Match xxx after headings
    if (t == "xxx" or t.startswith("xxx")) and i > 0:
        prev = doc.paragraphs[i-1]
        pvt = unicodedata.normalize("NFC", prev.text.strip())
        pvs = prev.style.name
        if pvs.startswith("Heading"):
            for k, v in data.items():
                if isinstance(v, str) and k in pvt:
                    p.clear(); p.add_run(v)
                    break

# Find section 4
body = doc.element.body
bh4 = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith("Heading") and "4." in p.text:
        for c in ["\u975e\u529f\u80fd","\u6027\u9700\u6c42"]:
            if c in p.text: bh4 = i; break
    if bh4: break

if bh4 is None:
    # Fallback
    for i, p in enumerate(doc.paragraphs):
        nt = unicodedata.normalize("NFC", p.text)
        if p.style.name.startswith("Heading") and "4." in nt and len(nt) <= 20:
            bh4 = i; break

print(f"Section 4 at: {bh4}")

# Insert modules 3.4-3.8
ref = list(body)[bh4]
pos = list(body).index(ref)
for sn in ["3.8","3.7","3.6","3.5","3.4"]:
    sec = data.get(sn)
    if not sec: continue
    elements = [
        mke(sec["interface"]),
        mke(f"{sn}.3 接口说明", h3_id),
        mke(sec["constraints"]),
        mke(f"{sn}.2 模块约束条件和输入输出", h3_id),
        mke(sec["overview"]),
        mke(f"{sn}.1 模块概述", h3_id),
        mke(f"{sn} {sec['title']}", h2_id),
    ]
    for elem in elements:
        body.insert(pos, elem)

# Fill sections 4,5
hmap = {"4.1":"4.1 性能需求","4.2":"4.2 安全和保密要求","4.3":"4.3 质量要求","5.1":"5.1错误信息","5.2":"5.2错误处理"}
for i, p in enumerate(doc.paragraphs):
    nt = unicodedata.normalize("NFC", p.text.strip())
    if not p.style.name.startswith("Heading"): continue
    for sk, sh in hmap.items():
        if (sh in nt or sk == nt) and sk in data and i+1 < len(doc.paragraphs):
            doc.paragraphs[i+1].clear()
            doc.paragraphs[i+1].add_run(data[sk])
            break

# Clear old template comments
for p in doc.paragraphs:
    t = p.text.strip()
    if "//" in t and len(t) > 20 and not p.style.name.startswith("Heading"):
        p.clear()

doc.save(OUT)
print(f"Done: {OUT}")

# Verify no xxx left
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t == "xxx" or t.startswith("xxx"):
        print(f"WARNING: xxx still at [{i}]")
