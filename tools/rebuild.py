# -*- coding: utf-8 -*-
"""Rebuild the functional spec from the original template in one pass."""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

SRC = r"D:\课设\功能规格说明模板.docx"
OUT = r"D:\project-89\功能规格说明_冷戰兵棋推演系统.docx"
doc = Document(SRC)

# Build a complete index of paragraphs with their text
# We"ll replace content paragraph by paragraph

# Map of heading text -> replacement content for the NEXT paragraph
replace_map = {}

# ===== 1.1 =====
replace_map["1.1产品说明"] = (
"本项目是一款面向冷戰背景的營級兵棋推演演示程序，基于Godot游戏引擎与C#语言开发。\n\n"
"系统模拟北约（NATO/蓝方）与华约（Warsaw Pact/红方）在德国富尔达缺口（Fulda Gap）地域的营级攻防对抗。"
"核心设计理念为系统驱动（System-Driven），采用热座模式（Hotseat）的双人交替回合制，涵盖地图机动、ZOC控制、"
"模块化战斗结算、后勤补给网络、疲劳度衰减、视野侦察、战线渲染与胜利得分追踪等完整兵棋系统闭环。\n\n"
"地图尺寸为50x30方形网格，包含平原、森林、半城镇、城镇四种自然地形与支线公路、高速公路两种基础设施图层。"
"每个营（Battalion）下辖若干连（Company），每个连下辖若干排（Platoon），排内容纳步兵班组与载具。"
"项目在极短的开发周期内，采用纯逻辑运算替代AI行为树，以JSON数据结构驱动单位模板与初始部署，"
"追求最精简的架构实现完整的兵棋系统体验。"
)

# ===== 1.2 =====
replace_map["1.2用户角色说明"] = (
"本系统为热座双人对战模式，用户角色分为两类：\n\n"
"（1）蓝方玩家（北约阵营）：在己方回合内，可指挥蓝方营级单位的移动与进攻。回合结束时系统自动执行后勤结算与疲劳恢复。\n\n"
"（2）红方玩家（华约阵营）：在己方回合内，可指挥红方营级单位的移动与进攻。回合结束时系统自动执行后勤结算与疲劳恢复。\n\n"
"系统本身不包含AI对手，所有决策由双方玩家交替在同一台设备上完成。"
"战斗部署阶段双方进入专门的战斗部署面板，拖拽单位至主攻、辅助、炮兵三个插槽后执行数值结算。"
)

# ===== 2.1 =====
replace_map["2.1软件环境"] = (
"操作系统：Windows 10/11（64位）\n"
"开发引擎：Godot 4.x（Mono版本，.NET 6/8 运行时）\n"
"编程语言：C#（Godot Mono 集成）\n"
"数据格式：JSON（单位模板、战斗序列、占领状态持久化）\n"
"第三方依赖：无额外NuGet包，全部使用Godot内置API与.NET标准库\n"
"渲染管线：Godot 3D场景树 + Control节点2D UI层叠\n"
"构建工具：MSBuild / Godot内置构建系统"
)

# ===== 2.2 =====
replace_map["2.2硬件环境"] = (
"CPU：x86-64 架构，主频 >= 2.0 GHz（推荐 2.5 GHz 以上）\n"
"内存：>= 4 GB RAM（推荐 8 GB）\n"
"显卡：支持 OpenGL ES 3.0 / Vulkan 的 GPU\n"
"存储空间：>= 200 MB 可用空间\n"
"显示器：分辨率 >= 1366 x 768\n"
"备注：本系统为单机热座模式，无需网络连接"
)

# ===== 3.1 =====
replace_map["3.1功能结构图"] = (
"系统由以下七大功能模块组成，GameSessionController 为核心调度器，接收玩家输入后委派至各子系统：\n\n"
"    (1) 战场网格与机动模块  -- GridMap / MovementResolver / ZOCManager / FrontlineResolver\n"
"    (2) 回合控制模块          -- TurnManager / GameFlowController / TurnPhaseRules / GameplayEventHub\n"
"    (3) 战斗计算模块          -- CombatResolver / CombatDeploymentPanel / EngagementResolver / CombatAutoDeployer\n"
"    (4) 后勤补给模块          -- SupplyManager / SupplyNetwork（基于Dijkstra加权图搜索）\n"
"    (5) 视野与渲染模块       -- VisionResolver / Grid3DRenderer / GameCamera / GameHud\n"
"    (6) 胜利判定模块          -- VictoryTracker（战损VP + 地理控制VP + CRT胜败判定）\n"
"    (7) 数据驱动层            -- UnitDatabase / UnitTemplate / BattalionFactory / JSON场景数据\n\n"
"模块间数据流：GameSessionController 接收输入 -> TurnManager 管理回合状态 -> "
"GridMap/MovementResolver 处理机动 -> CombatFlowController 触发战斗部署面板 -> "
"CombatResolver 完成数值结算。每回合结束时 SupplyManager 执行后勤结算，"
"VictoryTracker 更新得分，VisionResolver 刷新迷雾。"
)

# ===== Module sections data =====
module_sections = [
    # (old_heading_text_to_replace, new_heading_text, content_for_next_para)
    # We"ll find "3.2 xxx.." heading and replace it, then replace its sub-paragraphs
    
    # We need a different strategy - build the module sections by creating new paragraphs
    # and deleting old ones
]

# ===== Better approach: process sequentially =====
# Walk through paragraphs, and for each "xxx" placeholder after a known heading,
# replace with the appropriate content

# First, build the full list of replacements as (index, text) pairs
replacements = {}  # index -> replacement text

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not t:
        continue
    
    # Find the previous heading context
    prev_heading = ""
    for j in range(i-1, -1, -1):
        ps = doc.paragraphs[j].style.name
        if ps.startswith("Heading"):
            prev_heading = doc.paragraphs[j].text.strip()
            break
    
    # Check if this paragraph needs replacement
    for key, val in replace_map.items():
        if key == prev_heading and (t == "xxx" or "xxx" in t):
            replacements[i] = val
            break

# Apply replacements
for idx, text in replacements.items():
    doc.paragraphs[idx].clear()
    doc.paragraphs[idx].add_run(text)

print(f"Applied {len(replacements)} replacements")

# Now handle the "3.2 xxx" heading - rename it
# Also handle "3.3 xxx" heading
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if "3.2 xxx" in t:
        p.clear()
        p.add_run("3.2 战场网格与机动模块 (Battlefield Grid & Movement)")
    elif "3.3 xxx" in t:
        p.clear()
        p.add_run("3.3 战斗计算模块 (Combat Resolution)")

print("Headings renamed")

# Now we need to fill the sub-section content for 3.2 and 3.3
# Map of (section_sub, heading_text) -> content
sub_content = {
    # 3.2
    ("3.2.1", "模块概述"): (
"战场网格与机动模块负责管理50x30方形战场网格，包含地形系统、可通行性判定、移动路径查找、"
"移动动画以及ZOC（Zone of Control/控制区）动态计算。\n\n"
"地形系统采用"自然地形"与"基础设施"双图层独立结算逻辑。最终移动代价取两者最小值：\n"
"    Cost = min(Cost_terrain, Cost_infra)\n\n"
"地形类型：平原（AP消耗2）、森林（4）、半城镇（1）、城镇（1）。\n"
"基础设施：支线公路（AP消耗1）、高速公路（0.5）。\n\n"
"基础设施消耗低于地形消耗时自动生效，体现公路的战略机动优势。\n\n"
"ZOC机制：每个存活单位对其所在网格及周围8格（3x3切比雪夫距离为1的区域）投射ZOC。"
"任何单位无法进入敌方ZOC覆盖格。转弯裁剪（Corner Clipping）机制："
"斜向移动时，若两侧翼网格均为不可逾越障碍或被敌方占据，则移动判定被阻断。\n\n"
"核心类：GridMap / MovementResolver / ZOCManager / FrontlineResolver / TileData"
    ),
    ("3.2.2", "模块约束"): (
"约束条件：\n"
"（1）地图网格50x30，索引范围[0,49]x[0,29]，越界访问将被拒绝。\n"
"（2）移动AP容差使用EPSILON=0.05，判定公式为current_AP + EPSILON >= cost。\n"
"（3）移动消耗后AP保留1位小数，负微小值归零。\n"
"（4）不可通行地形不可进入，ZOC范围内敌方格子视为不可进入。\n"
"（5）BFS/泛洪算法从单位所在格向四周扩展。\n\n"
"输入：单位当前位置Vector2I、当前AP值、敌方位置集合、敌方ZOC集合、己方占位集合、网格地图。\n"
"输出：可达格子字典（格坐标->累计消耗AP）、最短路径（格子序列）、更新后的单位位置与剩余AP。"
    ),
    ("3.2.3", "接口说明"): (
"GridMap核心接口：\n"
"  TileData GetTile(Vector2I pos) -- 获取指定格子的地形数据\n"
"  bool IsInBounds(Vector2I pos) -- 检查坐标是否在地图范围内\n\n"
"MovementResolver核心接口：\n"
"  Dictionary<Vector2I, float> GetReachableTiles(pos, ap, isEnemyZOC, isOccupied, unit)\n"
"  List<Vector2I> FindPath(from, to, ap, isEnemyZOC, isOccupied, unit)\n\n"
"ZOCManager核心接口：\n"
"  HashSet<Vector2I> GetFactionZOC(IEnumerable<Vector2I> unitPositions)\n\n"
"FrontlineResolver核心接口：\n"
"  List<List<Vector2I>> ResolveFrontlineChains(int[,] occupationMap)"
    ),
    # 3.3
    ("3.3.1", "模块概述"): (
"战斗计算模块负责实现基于插槽（Slot-based）的模块化战斗数值结算系统。"
"核心采用黑盒概率演算--攻击方与防守方各自拥有4个固定席位的战斗插槽"
"（主攻营x2、辅助营x1、炮兵营x1），系统根据单位攻防属性、兵种能力标签、"
"地形修正、疲劳度系数、OOS补给惩罚等多重因子计算出优势分（Advantage Value V），"
"再查表（CRT）得出双方伤亡比例，最后通过加权随机抽选算法将伤害分配至营内各子单位。\n\n"
"战斗发起流程：\n"
"（1）攻击方点击敌方营（AP>=4，切比雪夫距离<=2即5x5区域），系统触发战斗。\n"
"（2）系统筛选区域内参战单位列表，弹出战斗部署面板。\n"
"（3）玩家拖拽己方营至主攻/辅助/炮兵三个插槽；敌方由贪心算法自动补位。\n"
"（4）执行结算：优势分公式 V = (A_basex(1+T_atk))/(D_basex(1+T_def)) + M_attr + E\n"
"（5）CRT映射表将V值映射为双方伤亡率（+1.5以上：攻3%防60%；...-1.5以下：攻60%防3%）。\n"
"（6）伤亡分配采用加权随机抽选循环：while(damagePool>0){计算总权重->随机抽取->HP-1->检测HP<30%MaxHp即消灭}\n\n"
"核心类：CombatResolver / CombatDeploymentPanel / EngagementResolver / CombatContext / "
"CombatAutoDeployer / CombatForce / CombatUtils"
    ),
    ("3.3.2", "模块约束"): (
"约束条件：\n"
"（1）战斗发起：攻击单位AP>=4、目标距离<=2（切比雪夫），超出拒绝发起。\n"
"（2）参战上限：每方最多4营（主攻x2+辅助x1+炮兵x1）。\n"
"（3）优势分V取值钳制[-10,10]；CRT七段：+1.5/+1.0/+0.5/0/-0.5/-1.0/-1.5。\n"
"（4）HP<30%MaxHp的子单位判定消灭（SurvivalState=0），不再参与后续计算。\n"
"（5）兵种修正：指挥网络缺失-2、无步兵-1、无侦察-1、反装甲惩罚等。\n"
"（6）疲劳度>=7时攻防系数x0.5，5-6时x0.9，>8时AP归零。\n"
"（7）OOS断联惩罚：turnsOOS=1时x0.8，>=2时x0.5。\n\n"
"输入：攻击/防守方营列表（List<Battalion>）、CombatContext（含阵营、地形加成、OOS回合数）、可选随机种子。\n"
"输出：CombatResolutionResult（优势分详情含各项修正明细、伤亡池点数、伤亡记录列表含IsDestroyed标记、疲劳增量）。"
    ),
    ("3.3.3", "接口说明"): (
"CombatResolver核心接口：\n"
"  CombatResolutionResult ResolveCombat(List<Battalion> attackers, List<Battalion> defenders, CombatContext ctx, ulong? seed)\n"
"  CombatResolutionResult PreviewCombat(...) -- 非破坏性预览\n"
"  AdvantageResult ComputeAdvantage(Battalion a, Battalion d, CombatContext ctx)\n\n"
"CombatDeploymentPanel：玩家拖拽交互、ExecuteCombat()触发结算\n"
"EngagementResolver：GetEngageableUnits(pos, friendlies) -- 5x5区域参战筛选\n"
"CombatUtils：IsInfantry/IsArtillery/HasCommandNetwork/HasAnyCapability等兵种判别"
    ),
}

# Apply sub-section content
# We need to match more carefully - find heading containing the section number
# and the sub-heading text, then the next para with "xxx"
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not t or "xxx" not in t:
        continue
    if i == 0:
        continue
    
    prev = doc.paragraphs[i-1]
    prev_t = prev.text.strip()
    prev_s = prev.style.name
    
    if not prev_s.startswith("Heading"):
        continue
    
    # Try to match against sub_content
    for (sec_num, sub_type), content in sub_content.items():
        if sec_num in prev_t and sub_type in prev_t:
            print(f"  Match: [{sec_num}] {sub_type} -> replace para {i}")
            doc.paragraphs[i].clear()
            doc.paragraphs[i].add_run(content)
            # Mark as done so we don"t double-match
            sub_content[(sec_num, sub_type)] = None
            break

# Apply to remaining unmatched
for (sec_num, sub_type), content in sub_content.items():
    if content is not None:
        print(f"  WARNING: Unmatched: {sec_num} {sub_type}")

# Now handle the "3.2.2" label specifically - it might be "3.2.2模块约束条件"
# since the search above used "模块约束" not "模块约束条件"
# Let me find any remaining "xxx" paragraphs near section headings

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t == "xxx" or t.startswith("xxx"):
        # Find the nearest heading before
        for j in range(i-1, max(i-5, -1), -1):
            pp = doc.paragraphs[j]
            if pp.style.name.startswith("Heading"):
                print(f"  Remaining xxx at [{i}], prev heading [{j}]: '{pp.text.strip()[:80]}'")
                break

print("Sub-sections processed")
doc.save(OUT)
print("Saved: " + OUT)
