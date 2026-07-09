# -*- coding: utf-8 -*-
from docx import Document
DOC = r"D:\project-89\tools\功能规格说明_已填写.docx"
doc = Document(DOC)
for i,p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t and ("4." in t or "5." in t or "非功能" in t or "错误" in t or "系统错误" in t):
        print(f"[{i}] [{p.style.name}] '{t[:120]}'")
    if i >= 104 and i <= 115:
        print(f"[{i}] [{p.style.name}] '{t[:120]}'")
print("---")
# Check what paragraphs are around 108-115
for i in range(104, min(120, len(doc.paragraphs))):
    t = doc.paragraphs[i].text.strip()
    print(f"[{i}] [{doc.paragraphs[i].style.name}] '{t[:150]}'")
