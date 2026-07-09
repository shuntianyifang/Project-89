import json
from docx import Document

doc = Document(r"D:\project-89\功能规格说明_冷戰兵棋推演系统.docx")

data = {}
for fn in ["spec_content.json","spec_content_2.json","spec_content_3.json","spec_content_4.json"]:
    with open(f"D:\\project-89\\tools\\{fn}", "r", encoding="utf-8") as f:
        data.update(json.load(f))

# Check each paragraph
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    s = p.style.name

    if (t == "xxx" or t.startswith("xxx")) and i > 0:
        prev = doc.paragraphs[i-1]
        pvt = prev.text.strip()
        pvs = prev.style.name
        
        if pvs.startswith("Heading"):
            matched = False
            for k, v in data.items():
                if isinstance(v, str) and k in pvt:
                    print(f"MATCH: [{i}] prev='{pvt[:60]}' key='{k}'")
                    matched = True
            if not matched:
                print(f"NO MATCH: [{i}] prev='{pvt[:60]}' text='{t}' style='{pvs}'")
