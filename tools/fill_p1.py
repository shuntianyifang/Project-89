# -*- coding: utf-8 -*-
from docx import Document
TEMPLATE = r"D:\课设\功能规格说明模板.docx"
OUTPUT = r"D:\project-89\tools\功能规格说明_已填写.docx"
doc = Document(TEMPLATE)
def np(txt):
    for i,p in enumerate(doc.paragraphs):
        if txt in p.text:
            return doc.paragraphs[i+1] if i+1<len(doc.paragraphs) else None
    return None
def sp(txt,text):
    p=np(txt)
    if p: p.clear(); p.add_run(text)

sp("1.1产品说明","本项目是一款面向冷戰背景的營級兵棋推演演示程序，基于Godot游戏引擎与C#语言开发。\n\n系统模拟北约（NATO/蓝方）与华约（Warsaw Pact/红方）在德国富尔达缺口（Fulda Gap）地域的营级攻防对抗。核心设计理念为系统驱动（System-Driven），采用热座模式（Hotseat）的双人交替回合制，涵盖地图机动、ZOC控制、模块化战斗结算、后勤补给网络、疲劳度衰减、视野侦察、战线渲染与胜利得分追踪等完整兵棋系统闭环。\n\n地图尺寸为50×30方形网格，包含平原、森林、半城镇、城镇四种自然地形与支线公路、高速公路两种基础设施图层。每个营（Battalion）下辖若干连（Company），每个连下辖若干排（Platoon），排内容纳步兵班组与载具。项目在极短的开发周期内，采用纯逻辑运算替代AI行为树，以JSON数据结构驱动单位模板与初始部署，追求最精简的架构实现完整的兵棋系统体验。")

sp("1.2用户角色说明","本系统为热座双人对战模式，用户角色分为两类：\n\n（1）蓝方玩家（北约阵营）：在己方回合内，可指挥蓝方营级单位的移动与进攻。回合结束时系统自动执行后勤结算与疲劳恢复。\n\n（2）红方玩家（华约阵营）：在己方回合内，可指挥红方营级单位的移动与进攻。回合结束时系统自动执行后勤结算与疲劳恢复。\n\n系统本身不包含AI对手，所有决策由双方玩家交替在同一台设备上完成。战斗部署阶段双方进入专门的战斗部署面板，拖拽单位至主攻、辅助、炮兵三个插槽后执行数值结算。")

sp("2.1软件环境","操作系统：Windows 10/11（64位）\n开发引擎：Godot 4.x（Mono版本，.NET 6/8 运行时）\n编程语言：C#（Godot Mono 集成）\n数据格式：JSON（单位模板、战斗序列、占领状态持久化）\n第三方依赖：无额外NuGet包，全部使用Godot内置API与.NET标准库\n渲染管线：Godot 3D场景树 + Control节点2D UI层叠\n构建工具：MSBuild / Godot内置构建系统")

sp("2.2硬件环境","CPU：x86-64 架构，主频 ≥ 2.0 GHz（推荐 2.5 GHz 以上）\n内存：≥ 4 GB RAM（推荐 8 GB）\n显卡：支持 OpenGL ES 3.0 / Vulkan 的 GPU\n存储空间：≥ 200 MB 可用空间\n显示器：分辨率 ≥ 1366 × 768\n备注：本系统为单机热座模式，无需网络连接")

sp("3.1功能结构图","系统由以下七大功能模块组成，GameSessionController 为核心调度器，接收玩家输入后委派至各子系统：\n\n    (1) 战场网格与机动模块  — GridMap / MovementResolver / ZOCManager / FrontlineResolver\n    (2) 回合控制模块          — TurnManager / GameFlowController / TurnPhaseRules / GameplayEventHub\n    (3) 战斗计算模块          — CombatResolver / CombatDeploymentPanel / EngagementResolver / CombatAutoDeployer\n    (4) 后勤补给模块          — SupplyManager / SupplyNetwork（基于Dijkstra加权图搜索）\n    (5) 视野与渲染模块       — VisionResolver / Grid3DRenderer / GameCamera / GameHud\n    (6) 胜利判定模块          — VictoryTracker（战损VP + 地理控制VP + CRT胜败判定）\n    (7) 数据驱动层            — UnitDatabase / UnitTemplate / BattalionFactory / JSON场景数据\n\n模块间数据流：GameSessionController 接收输入 → TurnManager 管理回合状态 → GridMap/MovementResolver 处理机动 → CombatFlowController 触发战斗部署面板 → CombatResolver 完成数值结算。每回合结束时 SupplyManager 执行后勤结算，VictoryTracker 更新得分，VisionResolver 刷新迷雾。")

# Phase 1 save
doc.save(OUTPUT)
print("Phase 1 OK: " + OUTPUT)
