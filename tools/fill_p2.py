# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT = r"D:\project-89\tools\功能规格说明_已填写.docx"
doc = Document(INPUT)

def fi(txt):
    for i,p in enumerate(doc.paragraphs):
        if txt in p.text: return i
    return None

def sp(txt, text):
    i = fi(txt)
    if i is not None and i+1 < len(doc.paragraphs):
        doc.paragraphs[i+1].clear()
        doc.paragraphs[i+1].add_run(text)

def set_h(section_num, title):
    i = fi(section_num + " xxx")
    if i is not None:
        doc.paragraphs[i].clear()
        doc.paragraphs[i].add_run(section_num + " " + title)

# ===== 3.2 =====
set_h("3.2", "战场网格与机动模块 (Battlefield Grid & Movement)")

sp("3.2.1模块概述", """战场网格与机动模块负责管理50x30方形战场网格，包含地形系统、可通行性判定、移动路径查找、移动动画以及ZOC（Zone of Control）动态计算。

地形系统采用"自然地形"与"基础设施"双图层独立结算逻辑。最终移动代价取两者最小值：Cost = min(Cost_terrain, Cost_infra)

地形类型：平原（AP消耗2）、森林（4）、半城镇（1）、城镇（1）。
基础设施：支线公路（AP消耗1）、高速公路（0.5）。

基础设施消耗低于地形消耗时自动生效，体现公路的战略机动优势。

ZOC机制：每个存活单位对其所在网格及周围8格（3x3切比雪夫距离为1的区域）投射ZOC。任何单位无法进入敌方ZOC覆盖格。转弯裁剪（Corner Clipping）机制：斜向移动时，若两侧翼网格均为不可逾越障碍或被敌方占据，则移动判定被阻断。

核心类：GridMap / MovementResolver / ZOCManager / FrontlineResolver / TileData""")

sp("3.2.2模块约束", """约束条件：
（1）地图网格50x30，索引范围[0,49]x[0,29]，越界访问将被拒绝。
（2）移动AP容差使用EPSILON=0.05，判定公式为current_AP + EPSILON >= cost。
（3）移动消耗后AP保留1位小数，负微小值归零。
（4）不可通行地形不可进入，ZOC范围内敌方格子视为不可进入。
（5）BFS/泛洪算法从单位所在格向四周扩展。

输入：单位当前位置Vector2I、当前AP值、敌方位置集合、敌方ZOC集合、己方占位集合、网格地图。
输出：可达格子字典（格坐标->累计消耗AP）、最短路径（格子序列）、更新后的单位位置与剩余AP。""")

sp("3.2.3接口说明", """GridMap核心接口：
  TileData GetTile(Vector2I pos)
  bool IsInBounds(Vector2I pos)

MovementResolver核心接口：
  Dictionary<Vector2I, float> GetReachableTiles(pos, ap, isEnemyZOC, isOccupied, unit)
  List<Vector2I> FindPath(from, to, ap, isEnemyZOC, isOccupied, unit)

ZOCManager核心接口：
  HashSet<Vector2I> GetFactionZOC(IEnumerable<Vector2I> unitPositions)

FrontlineResolver核心接口：
  List<List<Vector2I>> ResolveFrontlineChains(int[,] occupationMap)""")

# ===== 3.3 =====
set_h("3.3", "战斗计算模块 (Combat Resolution)")

# Find 3.3.1 in the 3.3 section (second occurrence)
cnt = 0
for i,p in enumerate(doc.paragraphs):
    if "3.3.1" in p.text and "模块概述" in p.text:
        cnt += 1
        if cnt == 2 and i+1 < len(doc.paragraphs):
            np = doc.paragraphs[i+1]
            if np.text.strip() == "xxx":
                np.clear()
                np.add_run("""战斗计算模块负责实现基于插槽（Slot-based）的模块化战斗数值结算系统。核心采用黑盒概率演算——攻击方与防守方各自拥有4个固定席位的战斗插槽（主攻营x2、辅助营x1、炮兵营x1），系统根据单位攻防属性、兵种能力标签、地形修正、疲劳度系数、OOS补给惩罚等多重因子计算出优势分V，再查表CRT得出双方伤亡比例，最后通过加权随机抽选算法将伤害分配至营内各子单位。

战斗发起流程：
（1）攻击方点击敌方营（AP>=4，切比雪夫距离<=2即5x5区域）。
（2）系统筛选区域内参战单位列表，弹出战斗部署面板。
（3）玩家拖拽己方营至主攻/辅助/炮兵三个插槽；敌方由贪心算法自动补位。
（4）执行结算：优势分公式 V = (A_base*(1+T_atk))/(D_base*(1+T_def)) + M_attr + E
（5）CRT映射表将V值映射为双方伤亡率（+1.5以上：攻3%防60%...-1.5以下：攻60%防3%）。
（6）伤亡分配采用加权随机抽选循环：while(damagePool>0) {计算总权重 -> 随机抽取 -> HP-1 -> 检测HP<30%MaxHp即消灭}

核心类：CombatResolver / CombatDeploymentPanel / EngagementResolver / CombatContext / CombatAutoDeployer / CombatForce / CombatUtils""")
            break

# 3.3.2
for i,p in enumerate(doc.paragraphs):
    if "3.3.2" in p.text and "模块约束" in p.text:
        np = doc.paragraphs[i+1] if i+1 < len(doc.paragraphs) else None
        if np and np.text.strip() == "xxx":
            np.clear()
            np.add_run("""约束条件：
（1）战斗发起：攻击单位AP>=4、目标距离<=2（切比雪夫），超出拒绝发起。
（2）参战上限：每方最多4营（主攻x2+辅助x1+炮兵x1）。
（3）优势分V取值钳制[-10,10]；CRT七段：+1.5/+1.0/+0.5/0/-0.5/-1.0/-1.5。
（4）HP<30%MaxHp的子单位判定消灭（SurvivalState=0）。
（5）兵种修正：指挥网络缺失-2、无步兵-1、无侦察-1、反装甲惩罚等。
（6）疲劳度>=7时攻防系数x0.5，5-6时x0.9，>8时AP归零。
（7）OOS断联惩罚：turnsOOS=1时x0.8，>=2时x0.5。

输入：攻击/防守方营列表、CombatContext（含阵营、地形加成、OOS回合数）、可选随机种子。
输出：CombatResolutionResult（优势分详情、伤亡池点数、伤亡记录列表含IsDestroyed标记、疲劳增量）。""")
        break

# 3.3.3
for i,p in enumerate(doc.paragraphs):
    if "3.3.3" in p.text and "接口说明" in p.text:
        np = doc.paragraphs[i+1] if i+1 < len(doc.paragraphs) else None
        if np:
            np.clear()
            np.add_run("""CombatResolver核心接口：
  CombatResolutionResult ResolveCombat(List<Battalion> attackers, List<Battalion> defenders, CombatContext ctx, ulong? seed)
  CombatResolutionResult PreviewCombat(...) -- 非破坏性预览
  AdvantageResult ComputeAdvantage(Battalion a, Battalion d, CombatContext ctx)

CombatDeploymentPanel：玩家拖拽交互、ExecuteCombat()触发结算
EngagementResolver：GetEngageableUnits(pos, friendlies) -- 5x5区域参战筛选
CombatUtils：IsInfantry/IsArtillery/HasCommandNetwork/HasAnyCapability等兵种判别""")
        break

print("Phase 2 OK")
doc.save(INPUT)
print("Saved")
