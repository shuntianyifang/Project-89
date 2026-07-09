"""Final spec: use paragraph index for 3.2/3.3 sections since Chinese chars differ."""
import json
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"D:\课设\功能规格说明模板.docx"
OUT = r"D:\project-89\功能规格说明_冷戰兵棋推演系统.docx"
BASE = r"D:\project-89\tools"

doc = Document(SRC)

# Load data
data = {}
for fn in ["spec_content.json","spec_content_2.json","spec_content_3.json","spec_content_4.json"]:
    with open(f"{BASE}\\{fn}", "r", encoding="utf-8") as f:
        data.update(json.load(f))

# Get style IDs
h2_id = h3_id = None
for s in doc.styles:
    if s.name == "Heading 2": h2_id = s.style_id
    elif s.name == "Heading 3": h3_id = s.style_id

def mke(text, sid=None):
    p = OxmlElement("w:p")
    if sid:
        pr = OxmlElement("w:pPr"); ps = OxmlElement("w:pStyle")
        ps.set(qn("w:val"), sid); pr.append(ps); p.append(pr)
    r = OxmlElement("w:r"); t = OxmlElement("w:t")
    t.text = text; t.set(qn("xml:space"), "preserve"); r.append(t); p.append(r)
    return p

# === STEP 1: Replace by direct index for known positions ===
# These are the paragraph indices we verified:
index_map = {
    # Section 1.1, 1.2, 2.1, 2.2
    55: data["1.1产品说明"],
    57: data["1.2用户角色说明"],
    61: data["2.1软件环境"],
    63: data["2.2硬件环境"],
    # Section 3.1
    66: data["3.1功能结构图"],
    # Section 3.2 main (xxx after heading)
    70: "战场网格与机动模块讨论战场地形、移动规则与ZOC控制系统，详见3.2.1-3.2.3各子章节。",
    # Section 3.2.1
    72: data["3.2.1模块介绍"],
    # Section 3.2.2
    74: data["3.2.2模块约束条件和输入输出"],
    # Section 3.2.3
    77: data["3.2.3接口说明"],
    # Section 3.3 main (xxx after heading)
    79: "战斗计算模块讨论基于插槽的模块化战斗结算规则，详见3.3.1-3.3.3各子章节。",
    # Section 3.3.1
    81: data["3.3.1模块介绍"],
    # Section 3.3.2
    83: data["3.3.2模块约束条件和输入输出"],
    # Section 3.3.3
    86: data["3.3.3接口说明"],
    # Section 4.1 - after heading at [90], content at [91]
    91: data["4.1"],
    # Section 4.2 - heading at [100], content at [101]
    101: data["4.2"],
    # Section 4.3 - after heading at [103]
    104: data["4.3"],
    # Section 5.1 - after heading at [108]
    109: data["5.1"],
    # Section 5.2 - after heading at [110]
    111: data["5.2"],
}

for idx, text in index_map.items():
    if idx < len(doc.paragraphs):
        doc.paragraphs[idx].clear()
        doc.paragraphs[idx].add_run(text)
    else:
        print(f"WARNING: index {idx} out of range")

# === STEP 2: Rename section headings ===
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not p.style.name.startswith("Heading"): continue
    if "xxx" in t:
        if "3.2" in t:
            p.clear(); p.add_run("3.2 战场网格与机动模块 (Battlefield Grid & Movement)")
        elif "3.3" in t:
            p.clear(); p.add_run("3.3 战斗计算模块 (Combat Resolution)")

# === STEP 3: Insert modules 3.4-3.8 before section 4 ===
body = doc.element.body
# Section 4 heading is at index 89 (original), but indices might shift
bh4 = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith("Heading") and "4." in p.text and len(p.text) <= 20:
        bh4 = i; break

if bh4:
    ref = list(body)[bh4]
    pos = list(body).index(ref)
    for sn in ["3.8","3.7","3.6","3.5","3.4"]:
        sec = data.get(sn)
        if not sec: continue
        for elem in [
            mke(sec["interface"]),
            mke(f"{sn}.3 接口说明", h3_id),
            mke(sec["constraints"]),
            mke(f"{sn}.2 模块约束条件和输入输出", h3_id),
            mke(sec["overview"]),
            mke(f"{sn}.1 模块概述", h3_id),
            mke(f"{sn} {sec['title']}", h2_id),
        ]:
            body.insert(pos, elem)

# === STEP 4: Clear old template comments ===
for p in doc.paragraphs:
    t = p.text.strip()
    if t.startswith("//") and len(t) > 20:
        p.clear()
    # Clear the old template filler after section 3.1
    if "功能流程" in t or "组织结构" in t or "逻辑视图" in t:
        p.clear()

# Clear old filler paragraphs between 91 and 100, 104-106
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    # These are the template's original guidance text paragraphs
    if "//" in t or ("说明" in t and len(t) < 50 and not p.style.name.startswith("Heading")):
        if any(kw in t for kw in ["系统安全性","例如","通常如下","可靠性","列出","处理相应"]):
            p.clear()

doc.save(OUT)
print(f"Done: {OUT}")

# Verify no xxx
remaining = 0
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t == "xxx" or t.startswith("xxx"):
        print(f"WARNING: xxx at [{i}]")
        remaining += 1
if remaining == 0:
    print("CLEAN: No xxx placeholders remaining")
