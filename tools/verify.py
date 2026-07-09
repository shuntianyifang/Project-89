from docx import Document
doc = Document(r"D:\project-89\功能规格说明_冷戰兵棋推演系统.docx")
heads = []
for i,p in enumerate(doc.paragraphs):
    s = p.style.name
    if s.startswith("Heading") or s.startswith("toc"):
        heads.append((s, p.text.strip()[:100]))
    # Also check for "xxx" unresolved placeholders
    if "xxx" in p.text.lower():
        heads.append(("*** UNFILLED ***", p.text.strip()[:100]))
for h in heads:
    print(f"[{h[0]}] {h[1]}")
print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
