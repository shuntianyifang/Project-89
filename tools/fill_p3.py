# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC = r"D:\project-89\tools\功能规格说明_已填写.docx"
doc = Document(DOC)

# ===== Insert 3.4-3.8 before section 4 heading (index 89) =====
body = doc.element.body

# Find heading style IDs
h2_id = None; h3_id = None
for s in doc.styles:
    if s.name == "Heading 2": h2_id = s.style_id
    elif s.name == "Heading 3": h3_id = s.style_id
if not h2_id: h2_id = "2"
if not h3_id: h3_id = "3"

def make_p(text, sid=None):
    p = OxmlElement("w:p")
    if sid:
        pPr = OxmlElement("w:pPr")
        ps = OxmlElement("w:pStyle")
        ps.set(qn("w:val"), sid)
        pPr.append(ps)
        p.append(pPr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    p.append(r)
    return p

ref = list(body)[89]  # Section 4 heading
pos = list(body).index(ref)

sections_data = [
    ("3.4", "后勤补给模块 (Supply System)",
     "后勤补给模块基于Dijkstra加权图搜索算法模拟物资沿交通网流动的物理特性，实现动态补给网络。核心类：SupplyManager（补给结算）、SupplyNetwork（Dijkstra寻路引擎）。\n\n补给源体系：绝对补给源（地图边缘出口，SP=36.0）、次级补给源（机场）。\n\n算法流程：（1）检测战略连通性：从绝对补给源开始正向蔓延（扣除地形AP阻力），接触枢纽时重新激活并重置SP=36.0。（2）检测次级连通性：若与老家断联但有机场，变更为Secondary_Supported，释放局部SP=18.0。（3）彻底瘫痪：若皆无，SP清零。\n\n补给阻断机制：敌方单位所在格阻力无穷大；ZOC额外增加15SP阻力；阵地门槛（剩余AP>=4的单位ZOC才产生阻断）。每回合结束时，补给范围外单位turnsOOS++，疲劳度累加。补给范围内单位按剩余AP梯度恢复疲劳（AP>=8恢复2，4<=AP<8恢复1），HP同步恢复（每恢复1点疲劳所有存活子单位HP+2至满血）。组织涣散单位（Fatigue>8）在补给范围内回到8。",
     "约束条件：（1）SP值范围[0,36]，36为满补给状态。（2）阻抗因素：地形AP消耗、敌方实体占据（无穷阻断）、ZOC阻断（+15）。（3）疲劳度安全性上限为20（规则阈值仍为8崩溃），OOS惩罚仅到2回合累积（>=2时疲劳+2/回合）。（4）疲劳恢复梯度基于剩余AP：>=8恢复2，[4,8)恢复1，<4无恢复。\n\n输入：我方单位列表、敌方占据位置集合、敌方ZOC集合、枢纽/机场位置、网格地图。\n输出：每个格SP值（float[,]）、各单位更新后的疲劳度、HP恢复量、turnsOOS状态。",
     "SupplyNetwork核心接口：float[,] ComputeSupplySP(map, faction, enemyOccupied, enemyZOC, enemyAP, hubs, airports)。\nSupplyManager核心接口：float[,] ComputeFactionSupplySP(faction, map, units, enemyOccupied, enemyZOC, hubs, airports)、void UpdateFactionEndTurn(...) -- 回合末批量更新补给状态。"),
    ("3.5", "回合控制模块 (Turn Control System)",
     "回合控制模块采用全局状态机管理双方交替行动与战斗流程的阶段切换。核心类：TurnManager / GameFlowController / TurnPhaseRules / TurnFlowController / GameplayEventHub。\n\n状态定义（GamePhase枚举）：StrategicMovement（大地图机动）、CombatDeployment_Attacker（攻击方部署）、CombatDeployment_Defender（防守方部署）、CombatResolution（数值结算）。\n\n双阵营FactionState枚举：Neutral(0)/Faction_Blue(1)/Faction_Red(2)。TurnManager管理CurrentFaction和CurrentPhase的完整状态机切换，记录当前交战上下文CombatContext。每2次阵营切换（双方各一次）计为1个完整回合，TurnNumber累进。\n\n输入拦截规则：玩家点击仅响应己方单位；非StrategicMovement阶段仅响应CombatContext内单位。Space键快捷结束回合，F6键循环补给覆盖层显示模式（关/己方/敌方/双方）。",
     "约束条件：（1）EndStrategicTurn仅允许在StrategicMovement阶段调用。（2）InitiateCombat仅允许在StrategicMovement阶段、且攻击单位为当前阵营、目标为敌方。（3）FinishAttackerDeployment/FinishDefenderDeployment各阶段严格顺序。（4）CancelCombat可在任意非StrategicMovement阶段调用。（5）ResetAP根据GetMaxAP()而非固定值，体现疲劳度对AP上限的影响。\n\n输入：玩家鼠标点击事件、键盘事件（Space/F6）。\n输出：阶段切换后状态、选中单位、可达格子、移动路径、战斗上下文初始化。",
     "TurnManager核心接口：void EndStrategicTurn()、void InitiateCombat(attacker,defender,ctx)、void FinishAttackerDeployment()/FinishDefenderDeployment(resolver)/CancelCombat()、void RegisterBattalion(Battalion b)。\nGameFlowController：管理选中状态（CurrentSelection/ReachableTiles）、移动动画回调链。\nGameplayEventHub：发布/订阅游戏事件。"),
    ("3.6", "视野与渲染模块 (Vision & Rendering)",
     "视野与渲染模块负责战场迷雾计算、3D网格渲染、2D UI HUD显示以及前线可视化。\n\n视野系统（VisionResolver）：瞎子基线（Base Vision=6）：无Recon标签营仅看到6格。建制内侦察激活（Standard Vision=8）：营内至少一个存活Recon实体，提升至8格。专业侦察营（Advanced Vision=12）：营种类判定为侦察营，直接12格。\n\n渲染系统（Grid3DRenderer）：纯色3D底座+2D纯白Billboard Sprite剪影。公路网Hub and Spoke涌现渲染。前线渲染：基于控制图遍历异色交界点，按连通段输出前沿链，3D ribbon mesh渲染。补给覆盖层：F6切换四种显示模式，热力色映射SP值。\n\n核心类：VisionResolver / Grid3DRenderer / GameCamera / GameHud / GameSceneBootstrapper",
     "约束条件：（1）视野仅在己方回合刷新，敌方回合仅显示己方侦察到的敌单位。（2）HUD面板响应窗口尺寸变化动态重排。（3）前线渲染链分段采样，避免过度细分影响帧率。\n\n输入：当前活跃阵营、所有单位列表、控制地图（occupationMap）、补给SP网格。\n输出：可视敌方单位列表、前线链段集合、HUD状态文本、补给覆盖层数据。",
     "VisionResolver核心接口：HashSet<Vector2I> UpdateGlobalVision(int currentFaction, List<(Battalion,Vector2I)> allUnits)。\nGrid3DRenderer核心接口：SetBlueUnits/SetRedUnits/SetActiveFaction/SetFrontlineChains/SetReachable/ShowPath/SetSupplyOverlayData、StartMoveAnimation(...)。\nGameHud核心接口：SetInfoText/SetStatusText/ShowOrgPanel/HideOrgPanel/UpdateVPPanel/SetCampaignCasualtyText。"),
    ("3.7", "胜利判定模块 (Victory System)",
     "胜利判定模块负责追踪战损VP与地理控制VP，基于PRD定义的CRT七级胜败区间给出战局评估。\n\n战损VP：每消灭一个敌方子单位，按该单位Cost值计入VP。RecordCombatResult统计每次战斗后被摧毁单位的Cost总和。\n\n地理控制VP：ScoreControlVP每回合结束时将控制格子数量直接计入VP。控制权判定基于ZOC：一格被一方ZOC覆盖且不被对方ZOC覆盖时由该方控制；双方ZOC重叠为中立区。\n\n战局评定（Victory CRT，R=BlueVP/RedVP钳制[0.1,10.0]）：R>=4.0决定性胜利 | 2.0<=R<4.0重大胜利 | 1.25<=R<2.0边缘胜利 | 0.8<=R<1.25血腥僵局 | 0.5<=R<0.8边缘失利 | 0.25<=R<0.5重大失利 | R<0.25一败涂地。\n\n核心类：VictoryTracker / VictoryAssessment / VictoryLevel",
     "约束条件：（1）VP比例钳制公式处理除零：双方均为0则R=1.0；一方为0则R=10（或0.1）。（2）子单位消灭判定基于IsDestroyed（HP<30%MaxHp），非HP归零。（3）控制图更新基于进入+ZOC规则：单位进入格立即覆写控制权；自方ZOC且不受敌方ZOC影响的格子覆写控制权；双方ZOC冲突格变为中立。\n\n输入：CombatResolutionResult、ControlMap、TurnNumber。\n输出：VictoryAssessment（含双方VP、控制格数、胜败等级、等级文案）。",
     "VictoryTracker核心接口：void RecordCombatResult(CombatResolutionResult result, int attackerFaction)、void ScoreControlVP()、VictoryAssessment Evaluate(int turnNumber)、int[,] UpdateOccupationFromEntryAndZOC(...)、string BuildCampaignCasualtySummary()。"),
    ("3.8", "数据驱动模块 (Data-Driven System)",
     "数据驱动模块实现逻辑代码与关卡数据全面分离的设计原则。所有单位属性、编成模板、战斗序列均通过外部JSON文件定义，无需修改代码即可调整兵力配置。\n\nJSON数据体系：(a)单位属性库(units_database.json)：定义每个单位ID的攻防数值、HP、兵种标签、暴露权重等。(b)编成模板(Templates/*.json)：定义各阵营营级编制（us/sov/frg/gdr各一）。(c)战斗序列(Scenarios/Fulda_Gap/oob_*.json)：指定各营的模板ID、初始坐标、结构覆盖、状态覆盖。(d)占领状态持久化(occupation_state.json)：保存/恢复控制图状态。\n\n核心类：UnitDatabase（单位属性注册表）、UnitTemplate（单位模板数据模型）、BattalionFactory（递归构建营-连-排-子单位树）、TemplateDatabase/TOEModels（编成模板解析器）。",
     "约束条件：（1）JSON文件采用UTF-8编码。（2）单位属性必须包含Attack/Defense/MaxHp/BaseWeight/Cost字段。（3）编成模板支持结构覆盖（add_nodes/remove_nodes）和状态覆盖（state_overrides），允许对同一模板自定义裁剪。（4）BattalionFactory按Company->Platoon->SubUnit三层递归构建，Composite（步+车）和Standard（纯单位）两种排类型。\n\n输入：场景标识、JSON文件路径。\n输出：完整的Battalion对象树（含SubUnitInstance列表）、控制图int[,]。",
     "UnitDatabase核心接口：static UnitTemplate GetTemplate(string unitId)、static void LoadFromJson(string jsonPath)。\nBattalionFactory核心接口：Battalion BuildFromTemplate(string templateId, Dictionary overrides)。\nFuldaGapScenario核心接口：void InitializeScenario() -- 加载全部数据并初始化战场。"),
]

# Insert in reverse order
for sn, title, overview, constraints, interface in reversed(sections_data):
    elements = [
        make_p(interface),
        make_p(sn + ".3 接口说明", h3_id),
        make_p(constraints),
        make_p(sn + ".2 模块约束条件和输入输出", h3_id),
        make_p(overview),
        make_p(sn + ".1 模块概述", h3_id),
        make_p(sn + " " + title, h2_id),
    ]
    for elem in elements:
        body.insert(pos, elem)

print("3.4-3.8 inserted")

# ===== Section 4 =====
# 4.1: heading at [90], content at [91]
p = doc.paragraphs[91]
p.clear()
p.add_run("""本系统为单机热座兵棋推演，性能需求主要集中在以下方面：
（1）并发用户：单机双人交替操作，非网络应用，无并发用户压力。
（2）响应时间：点击单位选择、格子移动、路径预览等操作响应时间不超过500ms。战斗结算（含CRT查表+伤害分配循环）单次耗时不超过100ms。
（3）帧率：3D场景渲染帧率维持30FPS以上（含50x30网格、单位Billboard、前线Ribbon Mesh、补给覆盖层四层渲染）。
（4）内存占用：运行时内存占用不超过500MB（含JSON数据解析、网格地图、单位树结构）。
（5）启动时间：从点击启动到主界面可交互不超过3秒。
（6）回合切换：回合末后勤结算（Dijkstra全图补给计算+疲劳恢复+控制图更新）不超过500ms。""")

# Clear the old template hints (paragraphs 92-99 are old template comments)
for i in range(92, 100):
    if i < len(doc.paragraphs):
        doc.paragraphs[i].clear()

# 4.2: heading at [100], content at [101]
p = doc.paragraphs[101]
p.clear()
p.add_run("""（1）数据安全：占领状态持久化采用user://路径（用户数据目录），不涉及系统目录写入。JSON配置文件采用只读方式加载，运行时不被修改。
（2）访问控制：热座模式设计上无网络通信，不存在远程攻击面。
（3）数据完整性：战斗结算使用可选随机种子（ulong? seed），支持回放验证。CombatResolutionResult包含完整修正明细，可追溯每场战斗的计算过程。
（4）用户确认：发动战斗前需玩家点击确认，未部署完成可取消（CancelCombat），避免误操作导致不可逆损失。""")

# Clear 4.2 old text
p = doc.paragraphs[102] if 102 < len(doc.paragraphs) else None
if p: p.clear()

# 4.3: heading at [103], content at [104,105,106]
p = doc.paragraphs[104]
p.clear()
p.add_run("""（1）可靠性：关键计算路径（战斗CRT、优势分计算、AP移动判定）有单元测试覆盖（Scripts/Tests/目录含9个测试文件）。浮点比较使用EPSILON容差，避免精度问题导致行为不确定。
（2）可维护性：代码按Scripts/目录分层（Core/Systems/Models/Factories/Data/Rendering/Scenarios），逻辑与数据分离。JSON驱动配置，修改兵种属性无需重新编译。
（3）可测试性：战斗结算支持随机种子参数，关键算法（SupplyNetwork、CombatResolver、MovementResolver）与Godot节点解耦，可在纯C#环境单元测试。
（4）可移植性：基于Godot跨平台引擎，代码使用.NET标准库+Godot内置API，无平台特定依赖。
（5）代码覆盖：现有测试覆盖战场网格、战斗结算、交战判定、视野系统、后勤补给、控制图编解码、编成覆盖等核心模块。""")

for i in [105, 106]:
    if i < len(doc.paragraphs):
        doc.paragraphs[i].clear()

print("Section 4 done")

# ===== Section 5 =====
# 5.1: heading at [108], content at [109]
p = doc.paragraphs[109]
p.clear()
p.add_run("""系统可能出现的错误信息及含义：
（1）"AP too low, need 4" -- 战斗发起时攻击方单位AP不足4，无法发动进攻。
（2）"Target too far, max 2" -- 战斗发起时目标距离超过2格（切比雪夫距离），超出交战范围。
（3）"Click to select" -- 正常状态提示，玩家未选择任何单位。
（4）"EndStrategicTurn只能在StrategicMovement阶段调用" -- 非机动阶段错误调用回合结束。
（5）"只能使用当前阵营的单位发起进攻" -- 尝试用敌方单位或非当前阵营单位发起攻击。
（6）"不能攻击己方单位" -- 攻击目标为友军时拒绝。
（7）"必须在AttackerDeployment/DefenderDeployment阶段调用" -- 战斗部署流程顺序错误。
（8）"Combat cancelled" -- 战斗部署中玩家取消操作。
（9）JSON文件缺失或格式错误时，FuldaGapScenario初始化抛出异常，提示具体文件路径和错误原因。""")

# 5.2: heading at [110], content at [111]
p = doc.paragraphs[111]
p.clear()
p.add_run("""错误处理策略：
（1）输入验证：所有玩家操作（单位选择、移动、攻击）在GameSessionController层经过状态检查（GameSessionRules.IsActionAllowed）和参数范围校验，非法操作被拒绝并通过HUD提示原因。
（2）防御性编程：优势分V钳制[-10,10]，VP比例钳制[0.1,10.0]，疲劳度钳制[0,20]，AP强制非负。除零操作在所有除法前检查分母是否为0（EPS=1e-6f）。
（3）阶段状态保护：TurnManager每阶段操作前检查CurrentPhase，不符合预期时抛出InvalidOperationException并由上层捕获。
（4）数据加载容错：JSON解析失败时FuldaGapScenario加载占领状态回退到默认全中立控制图，保证程序可运行。
（5）用户提示：所有不可操作的输入通过HUD InfoText显示原因（如"AP too low"），而非静默失败。
（6）调试支持：关键路径（战斗回调、VP更新、补给状态）输出GD.Print日志，方便运行时问题诊断。""")

print("Section 5 done")

doc.save(DOC)
print("All done: " + DOC)
