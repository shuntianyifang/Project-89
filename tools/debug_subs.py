import json
from docx import Document

doc = Document(r"D:\project-89\功能规格说明_冷戰兵棋推演系统.docx")

# Check paragraphs 70-87 specifically
for i in range(69, 90):
    p = doc.paragraphs[i]
    t = p.text.strip()
    s = p.style.name
    print(f"[{i}] [{s}] text='{t[:120]}'")

# Check what data keys we have
print("\n--- Data keys loaded ---")
data = {}
for fn in ["spec_content.json","spec_content_2.json","spec_content_3.json","spec_content_4.json"]:
    with open(f"D:\\project-89\\tools\\{fn}", "r", encoding="utf-8") as f:
        d = json.load(f)
        for k in d:
            v = d[k]
            if isinstance(v, str):
                print(f"  STR '{k}' = '{v[:50]}...'")
            elif isinstance(v, dict):
                print(f"  DICT '{k}'")

# Check if key is in heading
h71 = doc.paragraphs[71].text.strip()
print(f"\nHeading at [71]: '{h71}'")
for k in data:
    if isinstance(data[k], str) and "3.2" in k:
        print(f"  Check '{k}' in '{h71}': {k in h71}")
