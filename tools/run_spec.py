import json, sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"D:\课设\功能规格说明模板.docx"
OUT = r"D:\project-89\功能规格说明_冷戰兵棋推演系统.docx"
BASE = r"D:\project-89\tools"

doc = Document(SRC)

data = {}
for fn in ["spec_content.json","spec_content_2.json","spec_content_3.json","spec_content_4.json"]:
    with open(f"{BASE}\\{fn}", "r", encoding="utf-8") as f:
        data.update(json.load(f))

h2_id = h3_id = None
for s in doc.styles:
    if s.name == "Heading 2": h2_id = s.style_id
    elif s.name == "Heading 3": h3_id = s.style_id

def mke(text, sid=None):
    p = OxmlElement("w:p")
    if sid:
        pr = OxmlElement("w:pPr")
        ps = OxmlElement("w:pStyle")
        ps.set(qn("w:val"), sid)
        pr.append(ps)
        p.append(pr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    p.append(r)
    return p

body = doc.element.body
bh4 = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith("Heading") and "4." in p.text:
        v = p.text
        if any(c in v for c in ["\u975e\u529f\u80fd","\u6027\u9700\u6c42"]):
            bh4 = i
            break

if bh4 is None:
    for i, p in enumerate(doc.paragraphs):
        if "4." in p.text and ("\u975e" in p.text or "\u6027" in p.text):
            bh4 = i
            break

print(f"Section 4 at: {bh4}")

# Process paragraphs
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    st = p.style.name

    # Rename 3.2/3.3 xxx headings
    if st.startswith("Heading") and "xxx" in t:
        if "3.2" in t:
            p.clear(); p.add_run("3.2 \u6218\u573a\u7f51\u683c\u4e0e\u673a\u52a8\u6a21\u5757 (Battlefield Grid & Movement)")
        elif "3.3" in t:
            p.clear(); p.add_run("3.3 \u6218\u6597\u8ba1\u7b97\u6a21\u5757 (Combat Resolution)")

    # Replace xxx after headings
    if (t == "xxx" or t.startswith("xxx")) and i > 0:
        prev = doc.paragraphs[i-1]
        pvt = prev.text.strip()
        pvs = prev.style.name
        if pvs.startswith("Heading"):
            for k, v in data.items():
                if isinstance(v, str) and k in pvt:
                    p.clear(); p.add_run(v)
                    break

# Insert 3.4-3.8
ref = list(body)[bh4]
pos = list(body).index(ref)
for sn in ["3.8","3.7","3.6","3.5","3.4"]:
    sec = data.get(sn)
    if not sec: continue
    for elem in [
        mke(sec["interface"]),
        mke(f"{sn}.3 \u63a5\u53e3\u8bf4\u660e", h3_id),
        mke(sec["constraints"]),
        mke(f"{sn}.2 \u6a21\u5757\u7ea6\u675f\u6761\u4ef6\u548c\u8f93\u5165\u8f93\u51fa", h3_id),
        mke(sec["overview"]),
        mke(f"{sn}.1 \u6a21\u5757\u6982\u8ff0", h3_id),
        mke(f"{sn} {sec['title']}", h2_id),
    ]:
        body.insert(pos, elem)

# Fill 4 and 5
hmap = {"4.1":"4.1 \u6027\u80fd\u9700\u6c42","4.2":"4.2 \u5b89\u5168\u548c\u4fdd\u5bc6\u8981\u6c42","4.3":"4.3 \u8d28\u91cf\u8981\u6c42","5.1":"5.1\u9519\u8bef\u4fe1\u606f","5.2":"5.2\u9519\u8bef\u5904\u7406"}
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not p.style.name.startswith("Heading"): continue
    for sk, sh in hmap.items():
        if sh in t or sk == t:
            if sk in data and i+1 < len(doc.paragraphs):
                doc.paragraphs[i+1].clear()
                doc.paragraphs[i+1].add_run(data[sk])
            break

# Clear old template comments
for p in doc.paragraphs:
    t = p.text.strip()
    if "//" in t and len(t) > 20 and not p.style.name.startswith("Heading"):
        p.clear()

doc.save(OUT)
print(f"Done: {OUT}")
print(f"Paras: {len(doc.paragraphs)}")
